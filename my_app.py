"""
CV Builder Application

This script creates a CV (Curriculum Vitae) document using user input.
It uses the python-docx library to create a Word document and pyttsx3 for text-to-speech functionality.
"""
from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    """
    Use text-to-speech to speak the given text.

    Args:
        text (str): The text to be spoken.
    """
    pyttsx3.speak(text)
    
# Create a new Document
document = Document()

# Add profile picture
document.add_picture('me.jpg', width=Inches(2.0))

# Collect personal information
name = input('What is your name? ')
speak('Hello ' + name + ' how are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# Add 'About Me' section
document.add_heading('About Me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)


# Add work experience section
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('Enter start date: ')
to_date  = input('Enter to date: ')

p.add_run(company + '  ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('Enter start date: ')
        to_date  = input('Enter to date: ')

        p.add_run(company + '  ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break
    
# Add skills
document.add_heading('Skills')
skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break
    
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using GNM\'s CV generator'

# Save the document
document.save('cv.docx')